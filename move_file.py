import os
import pandas as pd
from docx import Document
import shutil  # Импортируем библиотеку для перемещения файлов

# Укажите имя папки
folder_name = "finel"

# Создание папки, если она не существует
os.makedirs(folder_name, exist_ok=True)

# Чтение CSV файла с новостями
df = pd.read_csv('generated_blog_texts.csv')  # Замените на ваше имя файла

# Создание нового документа Word
doc = Document()

# Перебор каждой строки в DataFrame
for index, row in df.iterrows():
    # Предполагаем, что текст находится в столбце "Текст блога"
    text = row['Текст блога']
    
    # Добавление текста в документ
    doc.add_paragraph(text)
    
    # Добавление пустой строки (опционально)
    doc.add_paragraph()

# Сохранение документа в папке "finel"
output_path = os.path.join(folder_name, 'news_document.docx')  # Замените на желаемое имя выходного файла
doc.save(output_path)

print(f"Документ сохранен как '{output_path}'")

# Перемещение файла generated_slogans.csv в папку finel
slogans_csv_path = 'generated_slogans.csv'  # Замените на путь к вашему файлу, если необходимо
destination_path = os.path.join(folder_name, 'generated_slogans.csv')

# Проверка, существует ли файл, и перемещение
if os.path.exists(slogans_csv_path):
    shutil.move(slogans_csv_path, destination_path)
    print(f"CSV файл перемещен как '{destination_path}'")
else:
    print(f"Файл '{slogans_csv_path}' не найден.")


